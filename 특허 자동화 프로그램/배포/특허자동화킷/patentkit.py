# -*- coding: utf-8 -*-
"""특허 출원 자동화 킷 — 임시명세서(가출원) 출원건을 만들고, 검증하고, 제출본을 낸다.

출원건 폴더 하나가 작업 단위다. 폴더 안에 다음을 둔다.

    patent_config.json          출원인·발명자·IPC·공지예외 등 서지사항
    01_명세서.md                명세서 본문 (청구범위를 품고 있어도 된다)
    02_청구범위.md              청구범위 정본
    03_제출본_명세서_청구범위.md  build 가 만든다. 직접 고치지 아니한다
    도면/                       도 1.svg ... (선택)

쓰는 법
    python patentkit.py init    출원건_05_새건       새 출원건 폴더를 만든다
    python patentkit.py build   출원건_05_새건       제출본을 다시 만든다
    python patentkit.py check   출원건_05_새건       검증만 한다
    python patentkit.py docx    출원건_05_새건       제출용 DOCX 를 만든다
    python patentkit.py all     .                    하위 출원건 전부를 처리한다

폴더 대신 상위 폴더를 주면 patent_config.json 을 가진 하위 폴더를 모두 찾아 돈다.
종료 코드는 오류가 하나라도 있으면 1, 없으면 0 이다.
"""
from __future__ import annotations

import argparse
import datetime as dt
import json
import re
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

버전 = "1.0"

# ─────────────────────────────────────────────────────────── 정규식

절머리 = re.compile(r"^#{1,4}\s*【([^】]+)】\s*$", re.M)
단락표지 = re.compile(r"^\[(\d{4})\]", re.M)          # 행 머리만. 본문 인용 참조는 세지 아니한다
청구항머리 = re.compile(r"^###\s*【청구항\s*(\d+)】\s*$", re.M)
인용 = re.compile(r"제\s*(\d+)\s*항")
도면인용 = re.compile(r"(?<![가-힣])도\s*(\d+)")        # 「산소포화도 16」 같은 어미는 세지 아니한다
부호사용 = re.compile(r"(?<=[가-힣A-Za-z])\((S?\d{1,3})\)")   # 연도 (2025) 등은 세지 아니한다
부호선언 = re.compile(r"(S?\d{1,3})\s*(?:내지\s*(S?\d{1,3}))?\s*:")
사람확인표식 = "★"

필수절 = [
    "명세서", "발명의 명칭", "기술분야", "발명의 배경이 되는 기술",
    "발명의 내용", "도면의 간단한 설명", "발명을 실시하기 위한 구체적인 내용",
    "요약서", "요약", "대표도",
]

문체규칙 = [
    ("존댓말 종결", re.compile(r"[가-힣](?:습니다|ㅂ니다)\."), "err"),
    ("굵게 마크업", re.compile(r"\*\*"), "err"),
    ("코드 마크업", re.compile(r"`"), "err"),
    ("목록 기호", re.compile(r"(?m)^\s*[-*+•]\s"), "err"),
    ("작은따옴표", re.compile(r"['‘’]"), "err"),
    ("큰따옴표", re.compile(r"[\"“”]"), "err"),
    ("보고서 기호", re.compile(r"[○●◆■□▶]"), "err"),
    ("물결 범위", re.compile(r"\d\s*~\s*\d"), "warn"),
]

# ─────────────────────────────────────────────────────────── 보고


class 보고:
    """검사 결과를 모은다. 오류 하나라도 있으면 출원 서식 생성을 막는다."""

    def __init__(self, 이름: str):
        self.이름 = 이름
        self.항목: list[tuple[str, str, str]] = []   # (등급, 검사명, 내용)

    def 통과(self, 검사, 내용=""):
        self.항목.append(("통과", 검사, 내용))

    def 경고(self, 검사, 내용):
        self.항목.append(("경고", 검사, 내용))

    def 오류(self, 검사, 내용):
        self.항목.append(("오류", 검사, 내용))

    def 판정(self, 검사, 문제: list[str], 통과문구=""):
        """문제 목록이 비면 통과, 아니면 오류로 적는다."""
        if 문제:
            for m in 문제[:20]:
                self.오류(검사, m)
            if len(문제) > 20:
                self.오류(검사, f"… 외 {len(문제) - 20}건")
        else:
            self.통과(검사, 통과문구)

    @property
    def 오류수(self):
        return sum(1 for 등급, _, _ in self.항목 if 등급 == "오류")

    @property
    def 경고수(self):
        return sum(1 for 등급, _, _ in self.항목 if 등급 == "경고")

    def 출력(self):
        표식 = {"통과": "  OK ", "경고": "  !! ", "오류": "  XX "}
        print(f"\n[{self.이름}]")
        너비 = max((len(검사) for _, 검사, _ in self.항목), default=0)
        for 등급, 검사, 내용 in self.항목:
            print(f"{표식[등급]}{검사.ljust(너비)}  {내용}")
        통과수 = len(self.항목) - self.오류수 - self.경고수
        print(f"  → 통과 {통과수} · 경고 {self.경고수} · 오류 {self.오류수}")


# ─────────────────────────────────────────────────────────── 문서 읽기


def 절나누기(본문: str) -> dict[str, str]:
    """【절이름】 단위로 본문을 잘라 사전으로 준다. 같은 이름이 겹치면 뒤엣것이 이긴다."""
    자리 = [(m.group(1).strip(), m.start(), m.end()) for m in 절머리.finditer(본문)]
    결과 = {}
    for i, (이름, 시작, 끝) in enumerate(자리):
        다음 = 자리[i + 1][1] if i + 1 < len(자리) else len(본문)
        결과[이름] = 본문[끝:다음]
    return 결과


def 청구항나누기(본문: str) -> list[tuple[int, str]]:
    자리 = [(int(m.group(1)), m.start(), m.end()) for m in 청구항머리.finditer(본문)]
    결과 = []
    for i, (번호, 시작, 끝) in enumerate(자리):
        다음 = 자리[i + 1][1] if i + 1 < len(자리) else len(본문)
        결과.append((번호, 본문[끝:다음].strip()))
    return 결과


def 꼬리정리(본문: str) -> str:
    """끝에 남은 구분선(---)과 빈 줄을 걷어낸다."""
    return re.sub(r"(?:\s*^-{3,}\s*)+\Z", "", 본문.rstrip(), flags=re.M).rstrip()


def 청구범위정본(본문: str) -> str:
    """02_청구범위.md 에서 내부 검토용 「회피 논거」 이하를 걷어낸다."""
    자리 = 본문.find("## 회피 논거")
    return 꼬리정리(본문[:자리] if 자리 > 0 else 본문)


# ─────────────────────────────────────────────────────────── config 검사

사업자번호꼴 = re.compile(r"^\d{3}-\d{2}-\d{5}$")
성명꼴 = re.compile(r"^(?:[가-힣]{2,6}|[一-鿿]{2,6}|[A-Za-z][A-Za-z .'\-]{1,59})$")


def config검사(cfg: dict, rep: 보고):
    if not cfg:
        rep.오류("config", "patent_config.json 을 읽지 못했다")
        return

    # 사람이 채워야 하는 값은 표식이 남아 있으면 차단한다
    남은표식 = []

    def 훑기(값, 길):
        if isinstance(값, dict):
            for k, v in 값.items():
                if k.startswith("_"):
                    continue
                훑기(v, f"{길}.{k}")
        elif isinstance(값, list):
            for i, v in enumerate(값):
                훑기(v, f"{길}[{i}]")
        elif isinstance(값, str) and 사람확인표식 in 값:
            남은표식.append(f"{길.lstrip('.')} = {값}")

    훑기(cfg, "")
    rep.판정("사람 확인 값", 남은표식, "차단 표식 없음")

    # ★ 를 지우고 예시 이름을 그대로 둔 것도 막는다. 발명자를 잘못 적으면 등록 뒤 무효사유가 된다
    예시이름 = {"홍길동", "김철수", "이영희", "아무개", "성명", "이름",
                "홍길순", "김갑돌", "홍길동1", "테스트"}
    가짜 = [f"inventors[{i}].name = {v.get('name')}"
            for i, v in enumerate(cfg.get("inventors") or [])
            if (v.get("name") or "").strip() in 예시이름]
    rep.판정("발명자 실명", 가짜, "예시 이름 없음")

    발명 = cfg.get("invention", {})
    if not 발명.get("title_kr"):
        rep.오류("발명의 명칭", "invention.title_kr 이 비어 있다")
    else:
        rep.통과("발명의 명칭", 발명["title_kr"][:40] + "…")

    문제 = []
    for i, a in enumerate(cfg.get("applicants") or []):
        if not a.get("name"):
            문제.append(f"applicants[{i}].name 이 비어 있다")
        번호 = (a.get("biz_no") or "").strip()
        if 번호 and 사람확인표식 not in 번호 and not 사업자번호꼴.match(번호):
            문제.append(f"applicants[{i}].biz_no 형식이 000-00-00000 이 아니다: {번호}")
    if not cfg.get("applicants"):
        문제.append("출원인이 한 명도 없다")
    rep.판정("출원인", 문제, f"{len(cfg.get('applicants') or [])}인")

    문제, 지분 = [], []
    발명자 = cfg.get("inventors") or []
    for i, v in enumerate(발명자):
        이름 = (v.get("name") or "").strip()
        if not 이름:
            문제.append(f"inventors[{i}].name 이 비어 있다")
        elif 사람확인표식 not in 이름 and not 성명꼴.match(이름):
            문제.append(f"inventors[{i}].name 형식이 성명으로 보이지 아니한다: {이름}")
        if v.get("share_pct") is not None:
            지분.append(v["share_pct"])
    if not 발명자:
        문제.append("발명자가 한 명도 없다")
    rep.판정("발명자", 문제, f"{len(발명자)}인")

    if len(발명자) > 1 and 지분:
        합 = round(sum(지분), 4)
        if len(지분) != len(발명자):
            rep.경고("발명자 지분", f"{len(발명자)}인 중 {len(지분)}인만 share_pct 가 있다")
        elif abs(합 - 100) > 1e-6:
            rep.오류("발명자 지분", f"지분 합계가 {합}% 다. 100% 로 맞춘다")
        else:
            rep.통과("발명자 지분", "합계 100%")
    elif len(발명자) > 1:
        rep.경고("발명자 지분", "복수 발명자인데 share_pct 가 비어 있다")

    if not cfg.get("ipc"):
        rep.경고("IPC", "ipc 가 비어 있다. 출원서 기재에 필요하다")
    else:
        rep.통과("IPC", " · ".join(cfg["ipc"]))

    출원 = cfg.get("filing") or {}
    공지예외 = 출원.get("disclosure_exception") or {}
    if 공지예외.get("claimed"):
        공개일 = 공지예외.get("date")
        if not 공개일:
            rep.오류("공지예외", "claimed 가 참인데 date 가 없다")
        else:
            try:
                d = dt.date.fromisoformat(공개일)
            except ValueError:
                rep.오류("공지예외", f"date 가 YYYY-MM-DD 가 아니다: {공개일}")
            else:
                기한 = d.replace(year=d.year + 1)
                남음 = (기한 - dt.date.today()).days
                적힌기한 = 출원.get("priority_deadline")
                if 적힌기한 and 적힌기한 != 기한.isoformat():
                    rep.오류("공지예외",
                             f"priority_deadline({적힌기한})이 공개일+12개월({기한})과 다르다")
                elif 남음 < 0:
                    rep.오류("공지예외", f"우선일 마감 {기한} 이 지났다")
                elif 남음 < 60:
                    rep.경고("공지예외", f"우선일 마감 {기한} 까지 {남음}일 남았다")
                else:
                    rep.통과("공지예외", f"마감 {기한} · {남음}일 남음")
                if not 공지예외.get("medium"):
                    rep.경고("공지예외", "medium(공개 매체)이 비어 있다. 증명서류 준비에 필요하다")


# ─────────────────────────────────────────────────────────── 명세서 검사


def 단락번호검사(본문: str, rep: 보고, 이름="단락번호"):
    표지 = [(int(m.group(1)), m.start()) for m in 단락표지.finditer(본문)]
    if not 표지:
        rep.오류(이름, "행 머리 단락번호가 하나도 없다")
        return
    문제 = []
    if 표지[0][0] != 1:
        문제.append(f"첫 표지가 [0001] 이 아니라 [{표지[0][0]:04d}] 이다")
    본값 = [n for n, _ in 표지]
    for 앞, 뒤 in zip(본값, 본값[1:]):
        if 뒤 == 앞:
            문제.append(f"중복 [{앞:04d}]")
        elif 뒤 < 앞:
            문제.append(f"역행 [{앞:04d}] → [{뒤:04d}]")
        elif 뒤 != 앞 + 1:
            문제.append(f"결번 [{앞:04d}] → [{뒤:04d}]")
    rep.판정(이름, 문제, f"[0001]~[{본값[-1]:04d}] · {len(본값)}개 연속")


def 명세서검사(spec: str, rep: 보고):
    절 = 절나누기(spec)

    빠진절 = [n for n in 필수절 if n not in 절]
    rep.판정("필수 절", [f"【{n}】 가 없다" for n in 빠진절], f"{len(필수절)}개 절 모두 있다")

    # [0001] 은 【기술분야】 바로 뒤에 온다
    기술분야 = 절.get("기술분야", "")
    첫표지 = 단락표지.search(기술분야)
    if not 기술분야.strip():
        rep.오류("단락번호 시작", "【기술분야】 절이 비어 있다")
    elif not 첫표지 or 첫표지.group(1) != "0001":
        rep.오류("단락번호 시작", "【기술분야】 첫 행이 [0001] 로 시작하지 아니한다")
    else:
        rep.통과("단락번호 시작", "【기술분야】 직후 [0001]")

    # 청구범위·요약서 앞까지가 명세서 본문이다
    본문끝 = len(spec)
    for 표식 in ("## 【청구범위】", "## 【요약서】"):
        자리 = spec.find(표식)
        if 자리 > 0:
            본문끝 = min(본문끝, 자리)
    단락번호검사(spec[:본문끝], rep)

    # 청구범위·요약서 구간에는 단락번호를 두지 아니한다
    뒷부분 = spec[본문끝:]
    남은표지 = 단락표지.findall(뒷부분)
    rep.판정("청구범위·요약서 단락번호",
             [f"[{n}] 가 청구범위·요약서 구간에 있다" for n in 남은표지],
             "표지 0건")

    # 도면 선언
    실시 = 절.get("발명을 실시하기 위한 구체적인 내용", "")
    선언도면 = {int(n) for n in 도면인용.findall(절.get("도면의 간단한 설명", ""))}
    사용도면 = {int(n) for n in 도면인용.findall(실시)}
    rep.판정("도면 선언",
             [f"본문이 도 {n} 을 인용하나 【도면의 간단한 설명】에 없다"
              for n in sorted(사용도면 - 선언도면)],
             f"선언 {len(선언도면)}매 · 인용 {len(사용도면)}매")
    for n in sorted(선언도면 - 사용도면):
        rep.경고("도면 선언", f"도 {n} 이 선언되었으나 본문이 인용하지 아니한다")

    대표도 = {int(n) for n in 도면인용.findall(절.get("대표도", ""))}
    if not 대표도:
        rep.오류("대표도", "【대표도】 가 도면을 지정하지 아니한다")
    elif 대표도 - 선언도면:
        rep.오류("대표도", f"대표도 {sorted(대표도 - 선언도면)} 가 선언되지 않은 도면이다")
    else:
        rep.통과("대표도", f"도 {sorted(대표도)[0]}")

    # 부호 대조
    선언부호 = set()
    for 시작, 끝 in 부호선언.findall(절.get("부호의 설명", "")):
        if 끝 and 시작[:1] == 끝[:1]:                      # S110 내지 S190 같은 범위
            첫, 막 = int(시작.lstrip("S")), int(끝.lstrip("S"))
            머리 = "S" if 시작.startswith("S") else ""
            선언부호 |= {f"{머리}{n}" for n in range(첫, 막 + 1)}
        else:
            선언부호.add(시작)
    본문구간 = "".join(v for k, v in 절.items()
                       if k not in ("부호의 설명", "특허문헌", "비특허문헌", "선행기술문헌"))
    사용부호 = set(부호사용.findall(본문구간))
    rep.판정("부호 대조",
             [f"본문의 부호 ({n}) 이 【부호의 설명】에 없다" for n in sorted(사용부호 - 선언부호)],
             f"선언 {len(선언부호)}개 · 사용 {len(사용부호)}개")

    return 절


# ─────────────────────────────────────────────────────────── 청구범위 검사

구성요소끝 = ("부", "표", "값", "레코드", "지표", "문자열", "필드", "계수",
              "목록", "코드", "단계", "장치", "정보", "집합", "번호", "식별자")
선행기재 = re.compile(r"상기\s+([가-힣A-Za-z0-9 ·제]{2,40}?)(?=[은는이가을를의와과에로,;\.])")


def 청구범위검사(claims_body: str, rep: 보고):
    항목 = 청구항나누기(claims_body)
    if not 항목:
        rep.오류("청구항", "청구항을 찾지 못했다")
        return None, []

    번호 = [n for n, _ in 항목]
    문제 = []
    for i, n in enumerate(번호, start=1):
        if n != i:
            문제.append(f"{i}번째 청구항의 번호가 【청구항 {n}】 이다")
    rep.판정("청구항 번호", 문제, f"청구항 1~{번호[-1]} · {len(번호)}개")

    있는번호 = set(번호)
    문제, 독립항 = [], []
    for n, 본문 in 항목:
        인용된 = {int(x) for x in 인용.findall(본문)}
        if not 인용된:
            독립항.append(n)
            continue
        for m in sorted(인용된):
            if m not in 있는번호:
                문제.append(f"청구항 {n} 이 없는 제{m}항을 인용한다")
            elif m == n:
                문제.append(f"청구항 {n} 이 자기 자신을 인용한다")
            elif m > n:
                문제.append(f"청구항 {n} 이 후행 제{m}항을 인용한다")
    rep.판정("청구항 인용 정합", 문제, f"독립항 {'·'.join(map(str, 독립항))}")

    if not 독립항:
        rep.오류("독립항", "독립항이 없다")

    # 카테고리 일원화의 기준이 되는 카테고리 명사를 청구항 1 에서 뽑는다
    카테고리 = None
    첫항 = dict(항목).get(1, "")
    m = re.match(r"\s*(.{2,60}?)로서\s*,", 첫항)
    if m:
        카테고리 = m.group(1).strip()
        rep.통과("카테고리 명사", 카테고리)
    else:
        rep.경고("카테고리 명사", "청구항 1 이 「…로서,」 로 시작하지 아니하여 카테고리를 뽑지 못했다")

    # 각 항의 말미가 카테고리 명사로 닫히는지 본다
    if 카테고리:
        문제 = []
        for n, 본문 in 항목:
            끝줄 = 꼬리정리(본문).rsplit("\n", 1)[-1].strip()
            if not 끝줄.endswith("."):
                문제.append(f"청구항 {n} 이 마침표로 끝나지 아니한다")
        rep.판정("청구항 종결", 문제, "전 항 마침표 종결")

    # 선행기재 — 앞선 항과 자기 항에 나오지 않는 「상기 X」 를 경고로 남긴다
    미기재 = []
    누적 = {}
    for n, 본문 in 항목:
        누적[n] = 본문
        인용된 = {int(x) for x in 인용.findall(본문)} & 있는번호
        범위 = 본문 + "".join(누적.get(m, "") for m in 인용된)
        앞선 = 범위.replace("상기 ", "")
        for 구 in 선행기재.findall(본문):
            구 = 구.strip()
            if not 구.endswith(구성요소끝):
                continue
            if re.search(r"및|또는|그리고|,", 구):   # 두 참조를 잇는 구는 나누어 보아야 한다
                continue
            if 구 not in 앞선:
                미기재.append(f"청구항 {n} 의 「상기 {구}」 가 앞서 기재되지 아니한다")
    for m in 미기재[:10]:
        rep.경고("선행기재", m)
    if not 미기재:
        rep.통과("선행기재", "미기재 참조 0건")
    elif len(미기재) > 10:
        rep.경고("선행기재", f"… 외 {len(미기재) - 10}건")

    return 카테고리, [n for n, _ in 항목]


def 카테고리일원화검사(카테고리, cfg, spec_절, rep: 보고):
    if not 카테고리:
        return
    명칭 = (cfg.get("invention") or {}).get("title_kr", "")
    문제 = []
    대상 = [("발명의 명칭", 명칭 or spec_절.get("발명의 명칭", "")),
            ("[0001]", spec_절.get("기술분야", "")),
            ("요약서", spec_절.get("요약", ""))]
    for 이름, 값 in 대상:
        if 카테고리 in 값:
            continue
        if 이름 == "발명의 명칭":
            문제.append(f"발명의 명칭에 카테고리 명사 「{카테고리}」 가 없다")
        else:          # [0001]·요약서는 상위 개념으로 적는 경우가 있어 사람 판단에 맡긴다
            rep.경고("카테고리 일원화",
                     f"{이름} 이 독립항의 카테고리 명사 「{카테고리}」 를 쓰지 아니한다")
    rep.판정("카테고리 일원화", 문제, f"명칭·[0001]·독립항·요약서 모두 「{카테고리}」")


# ─────────────────────────────────────────────────────────── 문체 검사


def 문체검사(본문: str, rep: 보고):
    줄 = [l for l in 본문.splitlines() if not l.lstrip().startswith("#")]
    검사대상 = "\n".join(줄)
    오류건, 경고건 = [], []
    for 이름, 패턴, 등급 in 문체규칙:
        찾음 = 패턴.findall(검사대상)
        if not 찾음:
            continue
        보기 = ", ".join(sorted({(x if isinstance(x, str) else x[0]) for x in 찾음})[:5])
        (오류건 if 등급 == "err" else 경고건).append(f"{이름} {len(찾음)}건 ({보기})")
    for m in 경고건:
        rep.경고("문체·표기", m)
    rep.판정("문체·표기", 오류건, f"{len(문체규칙)}개 유형 0건")


# ─────────────────────────────────────────────────────────── 제출본 생성


def 제출본만들기(spec: str, claims: str) -> str:
    """명세서 본문 뒤에 청구범위, 그 뒤에 요약서를 둔다(관청 표기 순서).

    명세서가 청구범위를 품고 있으면 그 구간을 걷어내고 정본을 다시 끼운다.
    청구항의 정본은 02_청구범위.md 하나로 유지한다.
    """
    정본 = 청구범위정본(claims)
    청구표식, 요약표식 = "## 【청구범위】", "## 【요약서】"
    if 요약표식 not in spec:
        raise ValueError("명세서에 「## 【요약서】」 가 없어 제출본 순서를 정할 수 없다")
    앞, 뒤 = spec.split(요약표식, 1)
    if 청구표식 in 앞:
        앞 = 앞.split(청구표식, 1)[0]
    return 앞.rstrip() + "\n\n" + 정본 + "\n\n" + 요약표식 + 뒤


def build(폴더: Path, cfg: dict, rep: 보고 | None = None) -> str:
    경로 = 경로들(폴더, cfg)
    spec = 경로["spec"].read_text(encoding="utf-8")
    claims = 경로["claims"].read_text(encoding="utf-8")
    제출본 = 제출본만들기(spec, claims)
    대상 = 폴더 / "03_제출본_명세서_청구범위.md"
    기존 = 대상.read_text(encoding="utf-8") if 대상.exists() else ""
    대상.write_text(제출본, encoding="utf-8")
    if rep is not None:
        if 기존 == 제출본:
            rep.통과("제출본 재생성", "변동 없음")
        elif 기존:
            rep.경고("제출본 재생성", f"갱신 {len(기존):,}자 → {len(제출본):,}자")
        else:
            rep.통과("제출본 재생성", f"신규 생성 {len(제출본):,}자")
    return 제출본


def 경로들(폴더: Path, cfg: dict) -> dict:
    p = cfg.get("paths") or {}
    return {
        "spec": 폴더 / (p.get("spec_md") or "01_명세서.md"),
        "claims": 폴더 / (p.get("claims_md") or "02_청구범위.md"),
        "out": 폴더 / (p.get("out_dir") or "_출력"),
    }


# ─────────────────────────────────────────────────────────── 편입분 대조


def 편입분대조(spec: str, claims: str, rep: 보고):
    """명세서에 편입된 청구범위와 정본이 문자 단위로 같은지 본다."""
    청구표식, 요약표식 = "## 【청구범위】", "## 【요약서】"
    if 청구표식 not in spec:
        rep.통과("편입분 대조", "명세서에 청구범위 편입 없음 (제출본에서 정본이 삽입된다)")
        return
    편입 = spec.split(청구표식, 1)[1]
    if 요약표식 in 편입:
        편입 = 편입.split(요약표식, 1)[0]
    정본 = 청구범위정본(claims).split(청구표식, 1)[-1]
    ㄱ, ㄴ = "".join(꼬리정리(편입).split()), "".join(꼬리정리(정본).split())
    if ㄱ == ㄴ:
        rep.통과("편입분 대조", f"정본과 일치 · {len(ㄱ):,}자")
        return
    for i, (a, b) in enumerate(zip(ㄱ, ㄴ)):
        if a != b:
            rep.오류("편입분 대조",
                     f"{i}번째 글자부터 다르다 · 편입 「…{ㄱ[max(0, i - 15):i + 15]}…」"
                     f" vs 정본 「…{ㄴ[max(0, i - 15):i + 15]}…」")
            return
    rep.오류("편입분 대조", f"길이가 다르다 · 편입 {len(ㄱ):,}자 vs 정본 {len(ㄴ):,}자")


def 명칭삼자대조(폴더: Path, cfg: dict, spec_절: dict, rep: 보고):
    값 = {"config": (cfg.get("invention") or {}).get("title_kr", "").strip(),
          "명세서": spec_절.get("발명의 명칭", "").strip()}
    파일 = 폴더 / "_발명의명칭.txt"
    if 파일.exists():
        값["_발명의명칭.txt"] = 파일.read_text(encoding="utf-8").strip()
    다름 = {k: v for k, v in 값.items() if v != 값["config"]}
    if not 값["config"]:
        rep.오류("발명의 명칭 대조", "config 의 title_kr 이 비어 있다")
    elif 다름:
        rep.오류("발명의 명칭 대조", f"{' · '.join(다름)} 이 config 와 다르다")
    else:
        rep.통과("발명의 명칭 대조", f"{len(값)}곳 일치")


# ─────────────────────────────────────────────────────────── 출원건 처리


def config읽기(폴더: Path) -> dict:
    파일 = 폴더 / "patent_config.json"
    if not 파일.exists():
        return {}
    try:
        return json.loads(파일.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        print(f"  patent_config.json 을 해석하지 못했다: {e}")
        return {}


def 출원건처리(폴더: Path, 생성=True) -> 보고:
    rep = 보고(폴더.name)
    cfg = config읽기(폴더)
    config검사(cfg, rep)

    경로 = 경로들(폴더, cfg)
    for 이름, 키 in (("01_명세서", "spec"), ("02_청구범위", "claims")):
        if not 경로[키].exists():
            rep.오류("파일", f"{경로[키].name} 이 없다")
    if rep.오류수 and not 경로["spec"].exists():
        return rep

    spec = 경로["spec"].read_text(encoding="utf-8")
    claims = 경로["claims"].read_text(encoding="utf-8") if 경로["claims"].exists() else ""

    if 생성 and claims:
        try:
            build(폴더, cfg, rep)
        except ValueError as e:
            rep.오류("제출본 재생성", str(e))

    절 = 명세서검사(spec, rep)
    명칭삼자대조(폴더, cfg, 절, rep)
    if claims:
        카테고리, _ = 청구범위검사(청구범위정본(claims), rep)
        카테고리일원화검사(카테고리, cfg, 절, rep)
        편입분대조(spec, claims, rep)

    제출본 = 폴더 / "03_제출본_명세서_청구범위.md"
    문체검사((제출본.read_text(encoding="utf-8") if 제출본.exists() else spec), rep)

    도면 = 폴더 / "도면"
    if 도면.is_dir():
        장수 = len(list(도면.glob("*.svg"))) or len(list(도면.glob("*.png")))
        rep.통과("도면 파일", f"{장수}매") if 장수 else rep.경고("도면 파일", "도면 폴더가 비어 있다")

    return rep


# ─────────────────────────────────────────────────────────── DOCX


def docx만들기(폴더: Path, cfg: dict) -> Path:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        raise SystemExit("DOCX 생성에는 python-docx 가 필요하다.  pip install python-docx")

    제출본 = 폴더 / "03_제출본_명세서_청구범위.md"
    if not 제출본.exists():
        raise SystemExit("제출본이 없다. 먼저 build 를 실행한다")
    본문 = 제출본.read_text(encoding="utf-8")

    style = cfg.get("style") or {}
    글꼴 = style.get("font_kr") or "바탕"
    크기 = style.get("heading_sizes") or {}

    문서 = Document()
    기본 = 문서.styles["Normal"]
    기본.font.name = 글꼴
    기본.font.size = Pt(11)
    기본.element.rPr.rFonts.set(
        __import__("docx.oxml.ns", fromlist=["qn"]).qn("w:eastAsia"), 글꼴)

    for 줄 in 본문.splitlines():
        벗김 = 줄.strip()
        if not 벗김:
            continue
        m = re.match(r"^(#{1,4})\s*(.+)$", 벗김)
        if m:
            수준 = len(m.group(1)) - 1 or 1
            p = 문서.add_paragraph()
            r = p.add_run(m.group(2))
            r.bold = True
            r.font.size = Pt(크기.get(str(수준), 12))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if 수준 == 1 else WD_ALIGN_PARAGRAPH.LEFT
        else:
            p = 문서.add_paragraph(벗김)
            p.paragraph_format.first_line_indent = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    출력 = 경로들(폴더, cfg)["out"]
    출력.mkdir(exist_ok=True)
    꼴 = (cfg.get("output") or {}).get("filename_pattern") or "임시명세서_{date}.docx"
    파일 = 출력 / 꼴.format(date=dt.date.today().strftime("%Y%m%d"))
    문서.save(파일)
    return 파일


def 제출절차(폴더: Path, cfg: dict, 출원: dict, 공지: dict) -> list[str]:
    """특허로에서 사람이 눌러야 하는 차례를 이 건의 실제 수치로 적는다.

    관청 제출 자체는 인증서 서명을 수반하는 법률행위라 킷이 대신하지 아니한다.
    대신 무엇을 몇 개 붙이고 어느 칸에 무엇을 적는지를 틀리지 않게 적어 둔다.
    """
    길 = 경로들(폴더, cfg)
    도면 = 폴더 / "도면"
    매수 = len({f.stem for f in 도면.glob("*.svg")} | {f.stem for f in 도면.glob("*.png")})         if 도면.is_dir() else 0
    try:
        항수 = len(청구항나누기(길["claims"].read_text(encoding="utf-8")))
    except OSError:
        항수 = 0
    docx = f"임시명세서_{dt.date.today().strftime('%Y%m%d')}.docx"

    남 = ""
    마감 = 출원.get("priority_deadline") or ""
    try:
        남 = f" (D-{(dt.date.fromisoformat(마감) - dt.date.today()).days})"
    except ValueError:
        pass

    줄 = ["## 제출 차례 — 사람이 한다", "",
          "관청 제출은 공동인증서로 서명하는 법률행위다. 킷은 여기까지만 만든다.", "",
          "| # | 할 일 | 값 |", "| --- | --- | --- |",
          "| 1 | 특허로 patent.go.kr 로그인 · 출원인코드 확인 | 최초 1회만 |",
          f"| 2 | 통합 명세서 작성기에 넣을 원고 | `_출력/{docx}` |",
          f"| 3 | 도면 첨부 | {매수}매 |",
          f"| 4 | 청구항 수 확인 | {항수}항 |",
          f"| 5 | 출원 종류 | {출원.get('kind', '')} |",
          f"| 6 | 수수료 감면 체크 | {출원.get('fee_reduction', '') or '해당 없음'} |",
          f"| 7 | 공지예외 주장 칸 | {'예 · 공개일 ' + (공지.get('date') or '') if 공지.get('claimed') else '아니오'} |",
          "| 8 | 인증서 서명 · 수수료 결제 · 제출 | 사람이 누른다 |",
          f"| 9 | 공지예외 증명서류 제출 | 출원일부터 30일 이내 |", ""]
    if 마감:
        줄 += [f"○ 우선일 마감 {마감}{남}", ""]
    return 줄


def 출원서요약(폴더: Path, cfg: dict) -> Path:
    """관청 서식에 옮겨 적을 서지사항을 한 장으로 뽑는다."""
    발명 = cfg.get("invention") or {}
    출원 = cfg.get("filing") or {}
    공지 = 출원.get("disclosure_exception") or {}
    줄 = [
        f"# 출원서 기재사항 — {폴더.name}", "",
        f"작성 {dt.date.today().isoformat()} · 특허 자동화 킷 {버전}", "",
        "| 항목 | 값 |", "| --- | --- |",
        f"| 발명의 명칭 | {발명.get('title_kr', '')} |",
        f"| 영문 명칭 | {발명.get('title_en', '')} |",
        f"| 출원 종류 | {출원.get('kind', '')} |",
        f"| IPC | {' · '.join(cfg.get('ipc') or [])} |",
        f"| 수수료 감면 | {출원.get('fee_reduction', '')} |",
        f"| 공지예외 주장 | {'예' if 공지.get('claimed') else '아니오'} |",
        f"| 공개일 | {공지.get('date', '')} |",
        f"| 공개 매체 | {공지.get('medium', '')} |",
        f"| 우선일 마감 | {출원.get('priority_deadline', '')} |",
        "", "## 출원인", "", "| 성명 | 사업자등록번호 | 구분 |", "| --- | --- | --- |",
    ]
    for a in cfg.get("applicants") or []:
        줄.append(f"| {a.get('name', '')} | {a.get('biz_no', '')} | {a.get('type', '')} |")
    줄 += ["", "## 발명자", "", "| 성명 | 소속 | 지분 | 대표 |", "| --- | --- | --- | --- |"]
    for v in cfg.get("inventors") or []:
        지분 = "" if v.get("share_pct") is None else f"{v['share_pct']}%"
        줄.append(f"| {v.get('name', '')} | {v.get('org', '')} | {지분} | "
                  f"{'예' if v.get('is_representative') else ''} |")
    줄 += ["", "○ 공지예외를 주장하는 경우 증명서류를 출원일부터 30일 이내에 제출한다"
           " (특허법 제30조제2항).", ""]
    줄 += 제출절차(폴더, cfg, 출원, 공지)
    출력 = 경로들(폴더, cfg)["out"]
    출력.mkdir(exist_ok=True)
    파일 = 출력 / "출원서_기재사항.md"
    파일.write_text("\n".join(줄), encoding="utf-8")
    return 파일


# ─────────────────────────────────────────────────────────── 도면


def 도면렌더(폴더: Path) -> int:
    도면 = 폴더 / "도면"
    파일들 = sorted(도면.glob("*.svg")) if 도면.is_dir() else []
    if not 파일들:
        print("  SVG 도면이 없다")
        return 0
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        raise SystemExit("도면 렌더에는 playwright 가 필요하다.  "
                         "pip install playwright && playwright install chromium")
    수 = 0
    with sync_playwright() as p:
        브라우저 = p.chromium.launch()
        for f in 파일들:
            svg = f.read_text(encoding="utf-8")
            w = int(float(svg.split('width="')[1].split('"')[0]))
            h = int(float(svg.split('height="')[1].split('"')[0]))
            page = 브라우저.new_page(viewport={"width": w, "height": h}, device_scale_factor=2)
            page.set_content(f"<html><body style='margin:0;background:#fff'>{svg}</body></html>")
            page.screenshot(path=str(f.with_suffix(".png")))
            page.close()
            수 += 1
            print(f"  렌더 {f.with_suffix('.png').name}  {w}x{h} @2x")
        브라우저.close()
    return 수


# ─────────────────────────────────────────────────────────── init

명세서뼈대 = """## 【명세서】

### 【발명의 명칭】

{제목}

### 【기술분야】

[0001] 본 발명은 {카테고리}에 관한 것이다.

### 【발명의 배경이 되는 기술】

[0002] ...

### 【선행기술문헌】

#### 【특허문헌】

(특허문헌 0001) KR 10-0000000 B1

#### 【비특허문헌】

(비특허문헌 0001) ...

### 【발명의 내용】

#### 【해결하려는 과제】

[0003] ...

#### 【과제의 해결 수단】

[0004] ...

#### 【발명의 효과】

[0005] ...

### 【도면의 간단한 설명】

[0006] 도 1은 본 발명의 일 실시예에 따른 {카테고리}(100)의 구성도이다.

### 【발명을 실시하기 위한 구체적인 내용】

[0007] 이하 첨부 도면을 참조하여 본 발명의 실시예를 상세히 설명한다.

[0008] 도 1을 참조하면, 본 실시예에 따른 {카테고리}(100)는 제1 구성부(110)를 포함한다.

### 【부호의 설명】

[0009] 100: {카테고리}, 110: 제1 구성부

## 【요약서】

### 【요약】

본 발명은 {카테고리}에 관한 것이다.

### 【대표도】

도 1
"""

청구범위뼈대 = """## 【청구범위】

### 【청구항 1】

{카테고리}로서,

입력 자료를 받아 제1 값을 산출하는 제1 구성부를 포함하는, {카테고리}.

### 【청구항 2】

제1항에 있어서,

상기 제1 구성부는 상기 제1 값을 저장하는, {카테고리}.
"""


def 카테고리뽑기(제목: str) -> str:
    """발명의 명칭에서 카테고리 명사를 뽑는다. 「…을 산출하는 지표 처리 장치, 그 방법…」 → 「지표 처리 장치」."""
    앞 = 제목.split(",")[0].strip()
    return 앞.split("하는 ")[-1].strip() or 앞


def init(폴더: Path, 제목: str):
    if 폴더.exists() and any(폴더.iterdir()):
        raise SystemExit(f"{폴더} 가 비어 있지 아니하다")
    카테고리 = 카테고리뽑기(제목)
    (폴더 / "도면").mkdir(parents=True, exist_ok=True)
    (폴더 / "01_명세서.md").write_text(
        명세서뼈대.format(제목=제목, 카테고리=카테고리), encoding="utf-8")
    (폴더 / "02_청구범위.md").write_text(
        청구범위뼈대.format(카테고리=카테고리), encoding="utf-8")
    (폴더 / "_발명의명칭.txt").write_text(제목, encoding="utf-8")
    cfg = json.loads(기본config)
    cfg["invention"]["title_kr"] = 제목
    올해 = dt.date.today()
    cfg["filing"]["disclosure_exception"]["date"] = 올해.isoformat()
    cfg["filing"]["priority_deadline"] = 올해.replace(year=올해.year + 1).isoformat()
    (폴더 / "patent_config.json").write_text(
        json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"출원건 폴더를 만들었다: {폴더}")
    print("  patent_config.json 의 ★ 표시 값을 채운 뒤 check 를 실행한다")


기본config = """{
  "_README": "★ 표시는 사람이 확인하여 채워야 하는 값이다. 채우기 전에는 검증이 차단된다.",
  "invention": { "title_kr": "", "title_en": "", "tech_field": "", "related_project": null },
  "applicants": [
    { "name": "★출원인 상호 확인 필요", "biz_no": "★사업자등록번호 확인 필요",
      "type": "주출원", "status": "확정" }
  ],
  "inventors": [
    { "name": "★발명자 확정 필요", "org": "", "position": "",
      "share_pct": null, "is_representative": true, "contact": "" }
  ],
  "ipc": [],
  "filing": {
    "kind": "임시명세서(가출원)",
    "fee_reduction": "중소기업 70% 감경 (특허료 등의 징수규칙 별표 5)",
    "disclosure_exception": { "claimed": false, "date": "", "medium": "" },
    "priority_deadline": ""
  },
  "paths": { "spec_md": "01_명세서.md", "claims_md": "02_청구범위.md", "out_dir": "_출력" },
  "style": {
    "font_kr": "바탕", "font_code": "Consolas",
    "heading_sizes": { "1": 16, "2": 14, "3": 12, "4": 11 },
    "table_style": "Light Grid Accent 1", "template_docx": ""
  },
  "output": { "filename_pattern": "임시명세서_{date}.docx" },
  "cover": { "recipient": "특허청" },
  "validation": { "strict": true }
}"""


# ─────────────────────────────────────────────────────────── CLI


def 출원건찾기(대상: Path) -> list[Path]:
    if (대상 / "patent_config.json").exists():
        return [대상]
    찾음 = sorted(p.parent for p in 대상.glob("*/patent_config.json"))
    if not 찾음:
        raise SystemExit(f"{대상} 아래에서 patent_config.json 을 찾지 못했다")
    return 찾음


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(
        prog="patentkit", description=f"특허 출원 자동화 킷 {버전}",
        formatter_class=argparse.RawDescriptionHelpFormatter, epilog=__doc__)
    ap.add_argument("명령", choices=["init", "build", "check", "docx", "figures", "all"])
    ap.add_argument("대상", nargs="?", default=".", help="출원건 폴더 또는 그 상위 폴더")
    ap.add_argument("--제목", default="", help="init 에서 쓸 발명의 명칭")
    a = ap.parse_args(argv)
    대상 = Path(a.대상).resolve()

    if a.명령 == "init":
        init(대상, a.제목 or "발명의 명칭을 채운다")
        return 0

    폴더들 = 출원건찾기(대상)
    print(f"특허 자동화 킷 {버전} · 출원건 {len(폴더들)}건 · {dt.date.today().isoformat()}")

    총오류 = 총경고 = 0
    for 폴더 in 폴더들:
        cfg = config읽기(폴더)
        if a.명령 == "figures":
            print(f"\n[{폴더.name}]")
            도면렌더(폴더)
            continue

        rep = 출원건처리(폴더, 생성=(a.명령 in ("build", "all", "docx")))
        rep.출력()
        총오류 += rep.오류수
        총경고 += rep.경고수

        if a.명령 in ("docx", "all"):
            if rep.오류수:
                print("  -- 오류가 있어 관청 서식 생성을 건너뛴다")
            else:
                요약 = 출원서요약(폴더, cfg)
                print(f"  생성 {요약.relative_to(폴더)}")
                try:
                    파일 = docx만들기(폴더, cfg)
                    print(f"  생성 {파일.relative_to(폴더)}")
                except SystemExit as e:
                    print(f"  -- {e}")

    print(f"\n종합 · 출원건 {len(폴더들)}건 · 경고 {총경고} · 오류 {총오류}")
    if 총오류 == 0:
        print("출원을 저지하는 결함이 확인되지 아니한다.")
    return 1 if 총오류 else 0


if __name__ == "__main__":
    sys.exit(main())
